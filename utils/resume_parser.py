# import PyPDF2
# import docx

# def parse_resume(file):
#     file_extension = file.name.split('.')[-1].lower()
    
#     if file_extension == 'pdf':
#         return parse_pdf(file)
#     elif file_extension == 'docx':
#         return parse_docx(file)
#     else:
#         raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

# def parse_pdf(file):
#     pdf_reader = PyPDF2.PdfReader(file)
#     text = ""
#     for page in pdf_reader.pages:
#         text += page.extract_text()
#     return text

# def parse_docx(file):
#     doc = docx.Document(file)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"
#     return text


import PyPDF2
import docx

def parse_resume(file):
    #file_extension = file.name.split('.')[-1].lower()
    if file and (file.filename.endswith('.pdf') ):
        return parse_pdf(file)
    elif file and (file.filename.endswith('.docx')):
        return parse_docx(file)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

def parse_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
        #print(text)
    return text

def parse_docx(file):
    doc = docx.Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text